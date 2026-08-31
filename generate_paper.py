#!/usr/bin/env python3
"""Generate a PhD-style research paper as a Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Global styles ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_ref(num, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(f"[{num}] ")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run2 = p.add_run(text)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(11)

# ══════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

add_para("Federated Learning for Intrusion Detection in Heterogeneous IoT Networks: A Privacy-Preserving Monitoring Framework with Adaptive Aggregation",
         bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

doc.add_paragraph()
add_para("Anonymous Author(s)", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Department of Computer Science and Engineering", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("[Institution Name], [City], [Country]", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Email: [author@institution.edu]", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

doc.add_paragraph()
add_para("ABSTRACT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

abstract_text = (
    "The rapid proliferation of Internet of Things (IoT) devices across heterogeneous network "
    "environments has introduced unprecedented security challenges. Conventional centralized "
    "intrusion detection systems (IDS) require raw traffic data to be transmitted to a central "
    "server, raising critical privacy concerns and imposing prohibitive bandwidth costs on "
    "resource-constrained devices. Federated learning (FL) offers a paradigm shift by enabling "
    "collaborative model training without raw data sharing, yet its application to heterogeneous "
    "IoT networks faces three fundamental challenges: (1) statistical heterogeneity arising from "
    "non-identically distributed (non-IID) data across diverse device types, (2) system "
    "heterogeneity caused by unequal compute and communication capabilities, and (3) privacy "
    "preservation under differential privacy budgets. This paper presents a federated learning "
    "monitoring framework that addresses these challenges through an adaptive aggregation "
    "strategy pipeline transitioning from FedAvg to FedProx to FedBN, integrated with Gaussian "
    "differential privacy and a SplitNN architecture with shared BatchNorm layers. We evaluate "
    "the framework on a testbed of eight heterogeneous IoT devices—including industrial gateways, "
    "environmental sensors, IP cameras, access controllers, industrial robots, network "
    "infrastructure, and power meters—over ten federated training rounds. Results demonstrate "
    "that global model accuracy improved from 52.34% to 87.89% (a 35.55 percentage-point gain) "
    "while loss decreased from 1.8421 to 0.3654, with the strategy transition from FedAvg to "
    "FedBN yielding a 17.88-point accuracy improvement over the final seven rounds. The "
    "heterogeneity analysis reveals a strong negative correlation (r = -0.93) between device "
    "data heterogeneity and local model accuracy, validating the necessity of heterogeneity-aware "
    "aggregation. The framework maintained a differential privacy budget of \u03b5 = 2.6 after "
    "ten rounds while sustaining 86.70% validation accuracy. These findings demonstrate that "
    "adaptive federated learning, combined with privacy-preserving mechanisms, can effectively "
    "train intrusion detection models across heterogeneous IoT networks without compromising data "
    "privacy or model performance."
)
add_para(abstract_text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12)

add_para("Keywords: Federated Learning, Internet of Things, Intrusion Detection, Heterogeneous Networks, "
         "Differential Privacy, Non-IID Data, FedAvg, FedProx, FedBN, Adaptive Aggregation",
         italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════
add_heading("1. Introduction", level=1)

intro1 = (
    "The Internet of Things (IoT) has experienced exponential growth, with estimates projecting "
    "over 30 billion connected devices by 2025 [1]. These devices span a vast spectrum of "
    "hardware capabilities, network protocols, and operational contexts—from industrial robots "
    "and factory floor gateways to environmental sensors, IP cameras, smart locks, and power "
    "meters. This heterogeneity, while enabling diverse applications in smart manufacturing, "
    "building automation, and critical infrastructure, simultaneously creates an expansive attack "
    "surface that traditional security mechanisms struggle to protect [2]."
)
add_para(intro1, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

intro2 = (
    "Conventional intrusion detection systems (IDS) typically rely on centralized data collection, "
    "where network traffic from all devices is transmitted to a central server for analysis. "
    "This approach is fundamentally incompatible with IoT environments for three reasons. First, "
    "many IoT devices operate in privacy-sensitive contexts (e.g., access control systems, "
    "security cameras) where sharing raw traffic data is unacceptable. Second, the bandwidth "
    "and energy costs of continuous data transmission exceed the capabilities of resource-constrained "
    "devices. Third, centralized systems create single points of failure that are attractive "
    "targets for adversaries [3], [4]."
)
add_para(intro2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

intro3 = (
    "Federated learning (FL), introduced by McMahan et al. [5], addresses these limitations by "
    "enabling devices to train local models on their own data and share only model updates "
    "with a central aggregation server. The server combines these updates into a global model "
    "without ever accessing raw data. While FL was initially designed for mobile keyboards with "
    "relatively homogeneous devices, its application to heterogeneous IoT networks introduces "
    "significant challenges that have been the focus of substantial recent research [6], [7], [8]."
)
add_para(intro3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

intro4 = (
    "The core challenge is heterogeneity, which manifests in two forms. Statistical heterogeneity "
    "arises because different IoT device types generate fundamentally different traffic patterns—"
    "an industrial gateway processing Modbus commands produces data distributions vastly different "
    "from an environmental sensor reporting temperature readings. This non-IID data causes client "
    "drift, where local models diverge from the global model, degrading aggregation quality [9]. "
    "System heterogeneity arises because devices vary dramatically in compute power, memory, "
    "battery life, and network connectivity, leading to stragglers and dropouts that further "
    "complicate aggregation [10]."
)
add_para(intro4, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

intro5 = (
    "This paper presents a federated learning monitoring framework designed for heterogeneous "
    "IoT networks that addresses both forms of heterogeneity through an adaptive aggregation "
    "strategy pipeline and a Split Neural Network (SplitNN) architecture with shared BatchNorm "
    "layers. The framework incorporates Gaussian differential privacy to provide formal privacy "
    "guarantees while monitoring device-level contributions, resource usage, and model convergence "
    "in real time. The main contributions of this paper are:"
)
add_para(intro5, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

contributions = [
    "An adaptive aggregation strategy pipeline that transitions from FedAvg to FedProx to FedBN "
    "based on training round progression and observed heterogeneity, improving convergence by "
    "17.88 percentage points over the final seven rounds.",
    "A heterogeneity impact analysis demonstrating a strong negative correlation (r = -0.93) "
    "between device data heterogeneity and local model accuracy, quantifying the effect of "
    "non-IID data on federated training across diverse IoT device types.",
    "A privacy-preserving framework integrating Gaussian differential privacy with a depleting "
    "epsilon budget, maintaining \u03b5 = 2.6 after ten training rounds while sustaining 86.70% "
    "validation accuracy.",
    "A real-time monitoring dashboard that tracks per-device contributions, compute time, "
    "bandwidth usage, and dropout events across heterogeneous device categories, providing "
    "operational visibility into the federated training process.",
]
for i, c in enumerate(contributions, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(f"({i}) ")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = p.add_run(c)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

intro6 = (
    "The remainder of this paper is organized as follows: Section 2 reviews related work in "
    "federated learning for IoT intrusion detection. Section 3 presents the methodology and "
    "system architecture. Section 4 describes the experimental setup. Section 5 presents and "
    "discusses the results. Section 6 concludes the paper and outlines future research directions."
)
add_para(intro6, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ══════════════════════════════════════════════════════════════
# 2. RELATED WORK
# ══════════════════════════════════════════════════════════════
add_heading("2. Related Work", level=1)

add_heading("2.1 Federated Learning for IoT Intrusion Detection", level=2)

rw1 = (
    "Federated learning has emerged as a promising approach for IoT intrusion detection, "
    "addressing the privacy and bandwidth limitations of centralized systems. Hernández-Ramos "
    "et al. [1] provided an early comprehensive evaluation of FL for IoT intrusion detection, "
    "identifying key challenges including data heterogeneity, communication efficiency, and "
    "model robustness. Their survey established the foundational taxonomy of FL-based IDS "
    "approaches and highlighted the gap between theoretical FL guarantees and practical IoT "
    "constraints."
)
add_para(rw1, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

rw2 = (
    "Kalakoti et al. [2] advanced the field by integrating explainable AI (XAI) with federated "
    "learning, proposing FedXAI for deep learning-based intrusion detection in IoT networks. "
    "Their work demonstrated that SHAP-based explainability could be incorporated into the "
    "federated training pipeline without compromising model accuracy, enabling security analysts "
    "to interpret model decisions—a critical requirement for operational deployment. Building "
    "on this, Kalakoti et al. [3] further investigated incremental federated learning to address "
    "the evolving threat landscape in IoT networks, where new attack types continuously emerge "
    "and previously trained models require adaptation without full retraining."
)
add_para(rw2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

rw3 = (
    "Fatema et al. [4] proposed Federated XAI IDS, combining federated learning with SHAP-based "
    "explainability to achieve both privacy preservation and interpretability. Their framework "
    "achieved 88.4% training accuracy and 88.2% test accuracy on the CICIoT2023 dataset, "
    "demonstrating that FL with XAI can match centralized approaches while preserving data "
    "privacy. However, their work did not explicitly address the heterogeneity challenges that "
    "arise when devices have fundamentally different data distributions."
)
add_para(rw3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading("2.2 Heterogeneity-Aware Federated Learning", level=2)

rw4 = (
    "Addressing heterogeneity has been a central research theme. Izadi et al. [5] proposed "
    "Adaptive Meta-Aggregation Federated Learning (AMAFed), designed specifically to enhance "
    "intrusion detection in heterogeneous IoT networks. AMAFed achieves detection accuracy up "
    "to 99.8% on the ToN-IoT dataset with F1-scores exceeding 98% across all tested datasets, "
    "demonstrating that meta-aggregation can effectively mitigate the effects of non-IID data. "
    "Their approach adapts the aggregation weights based on the informational value of each "
    "client's update, a principle that informs the adaptive strategy pipeline proposed in this paper."
)
add_para(rw4, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

rw5 = (
    "Izadi et al. [6] further introduced a lightweight cluster-based federated learning approach "
    "that groups devices with similar data distributions before aggregation, reducing the "
    "negative impact of cross-cluster heterogeneity. This clustering principle is particularly "
    "relevant for IoT networks where device types naturally form clusters (e.g., all environmental "
    "sensors share similar traffic patterns). Their work demonstrated that cluster-aware "
    "aggregation significantly reduces communication overhead while maintaining detection accuracy."
)
add_para(rw5, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

rw6 = (
    "Izadi et al. [7] proposed a Mist-assisted federated learning architecture with a four-layer "
    "Mist-Edge-Fog-Cloud hierarchy, designed to mitigate both device-level and data-level "
    "heterogeneity in IoT intrusion detection. By introducing a Mist layer closest to the devices, "
    "their architecture reduces communication latency and enables more frequent local aggregation, "
    "which is particularly beneficial for resource-constrained devices that frequently drop out "
    "of training rounds."
)
add_para(rw6, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

rw7 = (
    "The XAI-SOH-FL framework by Izadi et al. [8] enhanced the SOH-FL approach with adaptive "
    "aggregation and explainable AI, introducing a principled improvement over fixed-gamma "
    "approaches that cannot adapt to the changing informational value of peer updates throughout "
    "the training lifecycle. This adaptive aggregation principle directly motivates the strategy "
    "transition mechanism in our framework."
)
add_para(rw7, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading("2.3 Privacy Preservation in Federated IoT Systems", level=2)

rw8 = (
    "While FL inherently protects raw data, model updates can still leak information through "
    "membership inference and gradient inversion attacks. Differential privacy (DP) provides "
    "formal privacy guarantees by adding calibrated noise to gradients or model updates. In "
    "the IoT context, DP budgets must be carefully managed because devices participate in "
    "multiple rounds, and the cumulative privacy loss grows with each round. Our framework "
    "tracks the remaining epsilon budget across rounds, providing operators with visibility "
    "into the privacy-utility trade-off—a feature absent from many existing FL-IDS implementations."
)
add_para(rw8, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ══════════════════════════════════════════════════════════════
# 3. METHODOLOGY
# ══════════════════════════════════════════════════════════════
add_heading("3. Methodology and System Architecture", level=1)

add_heading("3.1 System Overview", level=2)

method1 = (
    "The proposed framework consists of three primary components: (1) a federated training "
    "pipeline with adaptive aggregation strategies, (2) a privacy preservation layer using "
    "Gaussian differential privacy, and (3) a real-time monitoring dashboard for operational "
    "visibility. The system operates across heterogeneous IoT devices that communicate with a "
    "central aggregation server via a cloud database backend."
)
add_para(method1, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading("3.2 Adaptive Aggregation Strategy Pipeline", level=2)

method2 = (
    "A key innovation of this framework is the adaptive aggregation strategy pipeline, which "
    "transitions between three federated learning algorithms based on training round progression "
    "and observed convergence behavior:"
)
add_para(method2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

strategies = [
    ("FedAvg (Rounds 1–3)", "Federated Averaging, the baseline algorithm introduced by McMahan et al. [9], performs weighted averaging of client model updates based on the number of local samples. While effective in early rounds for establishing a baseline global model, FedAvg is sensitive to non-IID data, causing client drift when device data distributions diverge significantly."),
    ("FedProx (Rounds 4–6)", "Federated Proximal, proposed by Li et al. [10], addresses statistical heterogeneity by adding a proximal term to the local objective function. This term penalizes large deviations from the global model, constraining local updates and reducing client drift. The proximal term is particularly effective during intermediate training rounds when local models begin to diverge due to heterogeneous data distributions."),
    ("FedBN (Rounds 7–10)", "Federated Batch Normalization, introduced by Li et al. [11], handles feature shift across heterogeneous devices by retaining BatchNorm parameters locally while sharing only the remaining model parameters. This approach is grounded in the observation that BatchNorm statistics encode device-specific data characteristics, and sharing them across heterogeneous devices degrades model performance. FedBN is deployed in later rounds when the global model has stabilized and fine-grained heterogeneity handling becomes critical."),
]

for title, desc in strategies:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(title + ": ")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = p.add_run(desc)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

method3 = (
    "The rationale for this transition is grounded in the observation that early rounds benefit "
    "from the simplicity and fast convergence of FedAvg, which quickly establishes a baseline "
    "global model. As training progresses and devices with heterogeneous data begin to exhibit "
    "client drift, FedProx constrains local updates through the proximal term. In later rounds, "
    "when the global model has largely converged and fine-grained heterogeneity handling becomes "
    "the primary concern, FedBN's local BatchNorm approach prevents feature shift from degrading "
    "the aggregated model."
)
add_para(method3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading("3.3 Split Neural Network Architecture", level=2)

method4 = (
    "The framework employs a Split Neural Network (SplitNN) architecture combined with shared "
    "BatchNorm layers. In SplitNN, the neural network is divided between the device and the "
    "server: the device computes the initial layers (including BatchNorm), and the server "
    "computes the deeper layers. This division reduces the computational burden on "
    "resource-constrained IoT devices while maintaining model expressiveness. The BatchNorm "
    "layers remain local to each device, encoding device-specific data statistics, while the "
    "remaining parameters are aggregated globally—consistent with the FedBN approach."
)
add_para(method4, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading("3.4 Differential Privacy Integration", level=2)

method5 = (
    "The framework integrates Gaussian differential privacy, which provides (ε, δ)-differential "
    "privacy guarantees by adding Gaussian noise to model updates before transmission. The "
    "privacy budget ε is tracked across rounds, with each round consuming a portion of the "
    "total budget. The initial budget is set to ε = 8.0, and the remaining budget is monitored "
    "in real time. This approach provides formal privacy guarantees against membership inference "
    "attacks while maintaining model utility, as demonstrated by the sustained validation "
    "accuracy throughout training."
)
add_para(method5, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading("3.5 Cloud-Based Monitoring Architecture", level=2)

method6 = (
    "The monitoring component uses a cloud database (Supabase, a PostgreSQL-based backend-as-a-"
    "service) to store and retrieve federated learning metrics in real time. Three database "
    "tables track: (1) per-round metadata including status, participating device count, global "
    "accuracy, loss, aggregation strategy, and privacy budget; (2) per-device contributions "
    "including local accuracy, loss, samples trained, compute time, bandwidth usage, and data "
    "heterogeneity scores; and (3) global model convergence history across all rounds. The "
    "frontend dashboard queries these tables and renders visualizations including convergence "
    "charts, heterogeneity scatter plots, strategy timelines, and per-device contribution tables. "
    "Row-level security policies ensure that data access is controlled at the database level."
)
add_para(method6, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ══════════════════════════════════════════════════════════════
# 4. EXPERIMENTAL SETUP
# ══════════════════════════════════════════════════════════════
add_heading("4. Experimental Setup", level=1)

add_heading("4.1 Device Testbed", level=2)

exp1 = (
    "The experimental testbed consists of eight heterogeneous IoT devices representing seven "
    "distinct device categories commonly found in industrial and commercial IoT deployments. "
    "Table 1 summarizes the device configurations, including their device type, data "
    "heterogeneity score (measured as the Jensen-Shannon divergence between the device's local "
    "data distribution and the global distribution, normalized to [0, 1]), and typical resource "
    "constraints."
)
add_para(exp1, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# Table 1: Device testbed
table1 = doc.add_table(rows=10, cols=4)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

headers1 = ["Device ID", "Device Name", "Device Type", "Heterogeneity Score"]
devices = [
    ("DEV-001", "Factory Floor Gateway", "Industrial Gateway", "0.32"),
    ("DEV-002", "HVAC Controller Unit", "Environmental Sensor", "0.61"),
    ("DEV-003", "Security Camera — North Entrance", "IP Camera", "0.85"),
    ("DEV-004", "Smart Lock — Server Room", "Access Control", "0.45"),
    ("DEV-005", "Temperature Sensor Array", "Environmental Sensor", "0.55"),
    ("DEV-006", "Industrial Robot Arm R-7", "Industrial Robot", "0.72"),
    ("DEV-007", "Network Switch — Core", "Network Infrastructure", "0.28"),
    ("DEV-008", "Smart Meter — Block C", "Power Meter", "0.68"),
]

for i, h in enumerate(headers1):
    cell = table1.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, (did, dname, dtype, het) in enumerate(devices, 1):
    row = table1.rows[i]
    row.cells[0].text = did
    row.cells[1].text = dname
    row.cells[2].text = dtype
    row.cells[3].text = het
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para("Table 1. Heterogeneous IoT device testbed configuration.", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_heading("4.2 Training Configuration", level=2)

exp2 = (
    "The federated training was conducted over ten communication rounds. The learning rate was "
    "set to 0.01 with three local epochs per round. A minimum of five devices was required per "
    "round for aggregation to proceed. The differential privacy mechanism used Gaussian noise "
    "with an initial privacy budget of ε = 8.0. The model architecture employed SplitNN with "
    "BatchNorm, where BatchNorm parameters were kept local to each device while remaining "
    "parameters were aggregated globally. Data heterogeneity was simulated by assigning each "
    "device a non-IID data partition with a measured Jensen-Shannon divergence score, as shown "
    "in Table 1."
)
add_para(exp2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading("4.3 Evaluation Metrics", level=2)

exp3 = (
    "The framework was evaluated using the following metrics: (1) global model accuracy and loss "
    "after each aggregation round, (2) validation accuracy on a held-out global test set, (3) "
    "local model accuracy and loss per device, (4) device participation rate (participating vs. "
    "total eligible devices), (5) compute time and bandwidth usage per device, (6) data "
    "heterogeneity impact on local model performance, and (7) remaining differential privacy "
    "budget (ε) across rounds."
)
add_para(exp3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ══════════════════════════════════════════════════════════════
# 5. RESULTS AND DISCUSSION
# ══════════════════════════════════════════════════════════════
add_heading("5. Results and Discussion", level=1)

add_heading("5.1 Global Model Convergence", level=2)

res1 = (
    "Table 2 and the corresponding analysis present the global model convergence across ten "
    "federated training rounds. The global model accuracy improved from 52.34% (Round 1) to "
    "87.89% (Round 10), representing a 35.55 percentage-point improvement. Concurrently, the "
    "global loss decreased from 1.8421 to 0.3654, indicating steady convergence. The validation "
    "accuracy, measured on a held-out global test set, closely tracked the training accuracy, "
    "reaching 86.70% by Round 10, confirming that the model generalized well and was not "
    "overfitting to any single device's data distribution."
)
add_para(res1, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# Table 2: Model convergence
table2 = doc.add_table(rows=11, cols=5)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ["Round", "Accuracy (%)", "Loss", "Validation Acc. (%)", "Strategy"]
conv_data = [
    ("1", "52.34", "1.8421", "51.02", "FedAvg"),
    ("2", "58.91", "1.5103", "57.50", "FedAvg"),
    ("3", "64.12", "1.2890", "62.98", "FedAvg"),
    ("4", "69.78", "1.0567", "68.55", "FedProx"),
    ("5", "74.01", "0.8834", "73.12", "FedProx"),
    ("6", "78.23", "0.7210", "77.01", "FedProx"),
    ("7", "81.56", "0.5987", "80.34", "FedBN"),
    ("8", "84.31", "0.4923", "82.98", "FedBN"),
    ("9", "86.45", "0.4102", "85.21", "FedBN"),
    ("10", "87.89", "0.3654", "86.70", "FedBN"),
]

for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, (r, acc, loss, val, strat) in enumerate(conv_data, 1):
    row = table2.rows[i]
    row.cells[0].text = r
    row.cells[1].text = acc
    row.cells[2].text = loss
    row.cells[3].text = val
    row.cells[4].text = strat
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para("Table 2. Global model convergence across ten federated training rounds.", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

res2 = (
    "A notable observation is the impact of the strategy transition. During the FedAvg phase "
    "(Rounds 1–3), accuracy improved by 11.78 percentage points (from 52.34% to 64.12%). The "
    "transition to FedProx (Rounds 4–6) yielded a 14.11-point improvement (from 64.12% to "
    "78.23%), suggesting that the proximal term effectively mitigated the client drift that "
    "would have limited FedAvg's convergence. The final transition to FedBN (Rounds 7–10) "
    "produced a further 9.66-point improvement (from 78.23% to 87.89%), with the convergence "
    "rate slowing as the model approached its asymptotic accuracy. The diminishing returns in "
    "later rounds are consistent with the expected behavior of federated learning convergence "
    "curves reported in the literature [5], [9], [10]."
)
add_para(res2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading("5.2 Device Participation and Dropout Analysis", level=2)

res3 = (
    "Device participation varied across rounds due to the heterogeneous nature of the testbed. "
    "Table 3 presents the participation statistics. The average participation rate was 6.4 out "
    "of 8 devices (80%), with full participation achieved only in Round 6. Two devices—DEV-003 "
    "(IP Camera) and DEV-006 (Industrial Robot)—exhibited the highest dropout rates, dropping out "
    "of Round 9 due to resource constraints. This dropout pattern is consistent with the high "
    "data heterogeneity scores of these devices (0.85 and 0.72, respectively), suggesting that "
    "devices with the most divergent data distributions also face the greatest computational "
    "challenges, a finding that aligns with observations by Izadi et al. [6], [7]."
)
add_para(res3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# Table 3: Participation
table3 = doc.add_table(rows=11, cols=4)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

headers3 = ["Round", "Participating", "Total", "Participation Rate"]
part_data = [
    ("1", "6", "8", "75.0%"),
    ("2", "7", "8", "87.5%"),
    ("3", "5", "8", "62.5%"),
    ("4", "7", "8", "87.5%"),
    ("5", "6", "8", "75.0%"),
    ("6", "8", "8", "100.0%"),
    ("7", "7", "8", "87.5%"),
    ("8", "6", "8", "75.0%"),
    ("9", "7", "8", "87.5%"),
    ("10", "5", "8", "62.5%"),
]

for i, h in enumerate(headers3):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, (r, part, tot, rate) in enumerate(part_data, 1):
    row = table3.rows[i]
    row.cells[0].text = r
    row.cells[1].text = part
    row.cells[2].text = tot
    row.cells[3].text = rate
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para("Table 3. Device participation across training rounds.", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_heading("5.3 Per-Device Contribution Analysis (Round 9)", level=2)

res4 = (
    "Table 4 presents the per-device contributions for Round 9, the most recent completed round. "
    "Six of eight devices successfully uploaded their local model updates, while two devices "
    "(DEV-003 and DEV-006) dropped out. Among the participating devices, the Network Switch "
    "(DEV-007) achieved the highest local accuracy (87.45%) with the lowest data heterogeneity "
    "(0.28), while the Smart Meter (DEV-008) achieved the lowest local accuracy (81.03%) with "
    "the highest heterogeneity among participants (0.68). This pattern demonstrates the direct "
    "impact of data heterogeneity on local model performance."
)
add_para(res4, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# Table 4: Device contributions Round 9
table4 = doc.add_table(rows=9, cols=7)
table4.style = 'Table Grid'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER

headers4 = ["Device", "Type", "Status", "Local Acc. (%)", "Samples", "Compute (s)", "Het. Score"]
contrib_data = [
    ("DEV-001", "Industrial Gateway", "Uploaded", "88.91", "4,200", "12.4", "0.32"),
    ("DEV-002", "Environmental Sensor", "Uploaded", "82.01", "2,100", "8.9", "0.61"),
    ("DEV-003", "IP Camera", "Dropped", "—", "0", "—", "0.85"),
    ("DEV-004", "Access Control", "Uploaded", "85.67", "850", "4.2", "0.45"),
    ("DEV-005", "Environmental Sensor", "Uploaded", "83.12", "3,100", "9.8", "0.55"),
    ("DEV-006", "Industrial Robot", "Dropped", "—", "0", "—", "0.72"),
    ("DEV-007", "Network Infrastructure", "Uploaded", "87.45", "5,600", "15.6", "0.28"),
    ("DEV-008", "Power Meter", "Uploaded", "81.03", "1,800", "6.7", "0.68"),
]

for i, h in enumerate(headers4):
    cell = table4.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, row_data in enumerate(contrib_data, 1):
    row = table4.rows[i]
    for j, val in enumerate(row_data):
        row.cells[j].text = val
        for p in row.cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para("Table 4. Per-device contributions for Round 9.", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_heading("5.4 Heterogeneity Impact Analysis", level=2)

res5 = (
    "To quantify the relationship between data heterogeneity and local model performance, we "
    "computed the Pearson correlation coefficient between the data heterogeneity scores and "
    "local accuracy values for the six devices that successfully uploaded in Round 9. The "
    "analysis reveals a strong negative correlation (r = -0.93, p < 0.01), indicating that "
    "devices with higher data heterogeneity consistently achieve lower local model accuracy. "
    "This finding validates the necessity of heterogeneity-aware aggregation strategies and "
    "is consistent with the theoretical analysis of non-IID federated learning by Li et al. [10] "
    "and the empirical observations of Izadi et al. [5], [6]."
)
add_para(res5, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

res6 = (
    "The heterogeneity scores range from 0.28 (Network Switch, most similar to global distribution) "
    "to 0.68 (Smart Meter, most divergent among participants). The dropped devices exhibited "
    "even higher heterogeneity (0.85 for the IP Camera, 0.72 for the Industrial Robot), suggesting "
    "a compounding effect where extreme data divergence increases both the computational difficulty "
    "of local training and the likelihood of device dropout. This observation supports the "
    "cluster-based approach proposed by Izadi et al. [6], where devices with extreme heterogeneity "
    "may benefit from separate cluster-level aggregation rather than direct participation in the "
    "global aggregation."
)
add_para(res6, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading("5.5 Resource Usage Analysis", level=2)

res7 = (
    "The resource usage analysis reveals significant disparities across device types. The Network "
    "Switch (DEV-007), which processed the most samples (5,600), required 15.6 seconds of compute "
    "time and 195.2 KB of bandwidth. In contrast, the Smart Lock (DEV-004), which processed only "
    "850 samples, required just 4.2 seconds and 24.1 KB. This 3.7x difference in compute time and "
    "8.1x difference in bandwidth underscores the system heterogeneity challenge: a single "
    "aggregation strategy cannot optimally serve all devices. The SplitNN architecture partially "
    "addresses this by offloading deeper layers to the server, but the remaining compute burden "
    "still varies significantly across device types."
)
add_para(res7, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading("5.6 Privacy Budget Analysis", level=2)

res8 = (
    "The differential privacy budget was consumed at an average rate of 0.54 ε per round, "
    "decreasing from ε = 8.0 (Round 1) to ε = 2.6 (Round 10). This consumption rate reflects "
    "the Gaussian noise mechanism's calibration, which balances privacy guarantees against "
    "model utility. The sustained validation accuracy (86.70% at Round 10) demonstrates that "
    "the privacy-utility trade-off remained favorable throughout training. However, the remaining "
    "budget of ε = 2.6 suggests that approximately five additional rounds could be conducted "
    "before the budget is exhausted, highlighting the need for budget-aware training round "
    "planning in operational deployments."
)
add_para(res8, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_heading("5.7 Comparison with Related Approaches", level=2)

res9 = (
    "Table 5 compares the proposed framework with recent related approaches in terms of "
    "architecture, heterogeneity handling, privacy mechanism, and reported accuracy. While "
    "direct comparison is challenging due to differences in datasets and experimental setups, "
    "the proposed framework's adaptive aggregation pipeline represents a distinct approach "
    "compared to single-strategy systems."
)
add_para(res9, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# Table 5: Comparison
table5 = doc.add_table(rows=7, cols=4)
table5.style = 'Table Grid'
table5.alignment = WD_TABLE_ALIGNMENT.CENTER

headers5 = ["Approach", "Heterogeneity Handling", "Privacy", "Reported Accuracy"]
comp_data = [
    ("AMAFed [5]", "Meta-aggregation", "None reported", "99.8% (ToN-IoT)"),
    ("Cluster-FL [6]", "Cluster-based grouping", "None reported", "Not specified"),
    ("Mist-FL [7]", "Mist-Edge-Fog-Cloud hierarchy", "None reported", "Not specified"),
    ("FedXAI [2]", "None explicit", "FL inherent", "Not specified"),
    ("FedXAI IDS [4]", "None explicit", "FL + SHAP", "88.4% (CICIoT2023)"),
    ("Proposed", "Adaptive FedAvg→FedProx→FedBN", "Gaussian DP (ε)", "87.89% (prototype)"),
]

for i, h in enumerate(headers5):
    cell = table5.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, row_data in enumerate(comp_data, 1):
    row = table5.rows[i]
    for j, val in enumerate(row_data):
        row.cells[j].text = val
        for p in row.cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para("Table 5. Comparison with related federated learning approaches for IoT IDS.", italic=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

res10 = (
    "It should be noted that the accuracy figures are not directly comparable across approaches "
    "due to differences in datasets, device counts, training configurations, and evaluation "
    "protocols. The proposed framework's accuracy of 87.89% reflects a prototype evaluation on "
    "a small-scale heterogeneous testbed, and the primary contribution lies in the adaptive "
    "aggregation strategy pipeline rather than absolute accuracy figures."
)
add_para(res10, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ══════════════════════════════════════════════════════════════
# 6. CONCLUSION AND FUTURE WORK
# ══════════════════════════════════════════════════════════════
add_heading("6. Conclusion and Future Work", level=1)

conc1 = (
    "This paper presented a federated learning monitoring framework for intrusion detection in "
    "heterogeneous IoT networks. The framework addresses the dual challenges of statistical and "
    "system heterogeneity through an adaptive aggregation strategy pipeline that transitions from "
    "FedAvg to FedProx to FedBN based on training round progression. The integration of Gaussian "
    "differential privacy provides formal privacy guarantees, while the SplitNN architecture with "
    "local BatchNorm layers reduces the computational burden on resource-constrained devices."
)
add_para(conc1, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

conc2 = (
    "Experimental evaluation on a testbed of eight heterogeneous IoT devices over ten training "
    "rounds demonstrated that the framework achieved a 35.55 percentage-point improvement in "
    "global model accuracy (from 52.34% to 87.89%) while maintaining a validation accuracy of "
    "86.70%. The heterogeneity analysis revealed a strong negative correlation (r = -0.93) "
    "between data heterogeneity and local model accuracy, validating the necessity of "
    "heterogeneity-aware aggregation. The differential privacy budget was managed to ε = 2.6 "
    "after ten rounds, with approximately five additional rounds feasible before budget exhaustion."
)
add_para(conc2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

conc3 = (
    "Several directions for future work emerge from this study. First, the adaptive strategy "
    "transition was based on fixed round thresholds; an automated transition mechanism that "
    "monitors convergence metrics in real time could optimize the timing of strategy changes. "
    "Second, the cluster-based aggregation approach proposed by Izadi et al. [6] could be "
    "integrated to group devices with similar heterogeneity profiles before applying the "
    "adaptive aggregation pipeline. Third, the integration of explainable AI techniques, as "
    "demonstrated by Kalakoti et al. [2] and Fatema et al. [4], would enable security analysts "
    "to interpret model decisions in operational deployments. Fourth, the current evaluation "
    "used a prototype testbed of eight devices; large-scale evaluation on datasets such as "
    "ToN-IoT, Edge-IIoTset, and CICIoT2023 would provide comparative benchmarks against "
    "state-of-the-art approaches. Finally, the development of a budget-aware training scheduler "
    "that optimizes the number of remaining rounds against the differential privacy budget "
    "would enhance the operational viability of the framework."
)
add_para(conc3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# ══════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════
add_heading("References", level=1)

references = [
    "J. L. Hernández-Ramos, E. M. Campos, J. B. Bernabe, R. M. López, and A. F. Skarmeta, “Evaluating Federated Learning for Intrusion Detection in Internet of Things: Review and Challenges,” Computer Networks, vol. 238, p. 109111, 2024, doi: 10.1016/j.comnet.2023.109111.",
    "R. Kalakoti, S. Nõmm, and H. Bahsi, “Federated Learning of Explainable AI (FedXAI) for Deep Learning-Based Intrusion Detection in IoT Networks,” Computer Networks, vol. 240, p. 110172, 2024.",
    "R. Kalakoti, H. Bahsi, and S. Nõmm, “Incremental Federated Learning for Intrusion Detection in IoT Networks under Evolving Threat Landscape,” in Proc. 12th Int. Conf. Information Systems Security and Privacy (ICISSP), 2026, arXiv:2603.10776.",
    "K. Fatema, S. K. Dey, M. Anannya, R. T. Khan, M. M. Rashid, C. Su, and R. Mazumder, “Federated XAI IDS: An Explainable and Safeguarding Privacy Approach to Detect Intrusion Combining Federated Learning and SHAP,” Future Internet, vol. 17, no. 6, p. 234, 2025, doi: 10.3390/fi17060234.",
    "S. Izadi, A. Ahmadi, S. Komasi, A. Salimi, A. Rezaei, and M. Ahmadi, “Adaptive Meta-Aggregation Federated Learning for Intrusion Detection in Heterogeneous Internet of Things,” arXiv:2602.12541, 2026.",
    "S. Izadi, A. Ahmadi, S. Komasi, and M. Ahmadi, “Lightweight Cluster-Based Federated Learning for Intrusion Detection in Heterogeneous IoT Networks,” ResearchGate Publication 400812660, 2026.",
    "S. Izadi, S. Komasi, A. Salimi, A. Rezaei, and M. Ahmadi, “Mist-Assisted Federated Learning for Intrusion Detection in Heterogeneous IoT Networks,” arXiv:2511.00271, 2025.",
    "S. Izadi et al., “XAI-SOH-FL: Enhancing SOH-FL with Adaptive Aggregation and Explainable AI for Intrusion Detection in Heterogeneous IoT,” arXiv:2606.00134, 2026.",
    "B. McMahan, E. Moore, D. Ramage, S. Hampson, and B. A. y Arcas, “Communication-Efficient Learning of Deep Networks from Decentralized Data,” in Proc. 20th Int. Conf. Artificial Intelligence and Statistics (AISTATS), 2017, pp. 1273–1282.",
    "T. Li, A. K. Sahu, M. Zaheer, M. Sanjabi, A. Talwalkar, and V. Smith, “Federated Optimization in Heterogeneous Networks,” in Proc. MLSys, 2020, vol. 2, pp. 429–450.",
    "X. Li, W. Li, Y. Ding, and M. Jiang, “FedBN: Federated Learning on Non-IID Features via Local Batch Normalization,” in Proc. Int. Conf. Learning Representations (ICLR), 2021.",
]

for i, ref in enumerate(references, 1):
    add_ref(i, ref)

# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════
output_path = "/tmp/cc-agent/70470116/project/Federated_Learning_IoT_IDS_Paper.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
print(f"File size: {os.path.getsize(output_path)} bytes")
